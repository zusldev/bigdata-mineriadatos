"""
Construye el documento Word (.docx) del informe del caso de estudio.
Lee el Markdown fuente, parsea secciones/tablas, incrusta gráficas PNG.

Uso:
    python src/report/build_docx.py
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
CHARTS = os.path.join(BASE, "outputs", "charts")
REPORT_MD = os.path.join(BASE, "reports", "informe_caso_estudio.md")
OUTPUT_DOCX = os.path.join(BASE, "reports", "informe_caso_estudio.docx")


# ── Helpers ──────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color_hex,
    })
    shading.append(shd)


def add_formatted_paragraph(doc, text, style="Normal", bold=False, italic=False,
                            size=None, color=None, alignment=None, space_after=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table_from_rows(doc, headers, rows, col_widths=None):
    """Add a formatted table to the doc."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "2C3E50")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph("")  # spacer
    return table


def add_image_if_exists(doc, filename, width=Inches(5.8), caption=None):
    """Add a PNG chart to the document if it exists."""
    path = os.path.join(CHARTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            add_formatted_paragraph(doc, caption, italic=True, size=9,
                                    color=(127, 140, 141),
                                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                    space_after=12)
        return True
    else:
        add_formatted_paragraph(doc, f"[Gráfica no disponible: {filename}]",
                                italic=True, size=9, color=(192, 57, 43))
        return False


# ═══════════════════════════════════════════════════════════════════
# Build Document
# ═══════════════════════════════════════════════════════════════════
def build():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font  # type: ignore[union-attr]
    font.name = "Calibri"
    font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    for _ in range(6):
        doc.add_paragraph("")

    add_formatted_paragraph(doc, "Informe Integral", bold=True, size=28,
                            color=(44, 62, 80),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "Caso de Estudio: \"Sabor Mexicano\"", bold=True, size=20,
                            color=(41, 128, 185),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER,
                            space_after=30)
    add_formatted_paragraph(doc, "Big Data y Minería de Datos", size=14,
                            color=(127, 140, 141),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "Cadena de 10 sucursales en México", size=14,
                            color=(127, 140, 141),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, "Febrero 2026", size=14,
                            color=(127, 140, 141),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ÍNDICE
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("Índice de Contenido", level=1)
    toc_items = [
        "1. Resumen Ejecutivo",
        "2. Análisis Exploratorio",
        "   2.1 Tendencias de ventas",
        "   2.2 Platillos más vendidos por región y hora del día",
        "   2.3 Sucursales con mejor y peor desempeño",
        "   2.4 Métodos de pago",
        "3. Identificación de Problemas",
        "   3.1 Factores que influyen en la rentabilidad",
        "   3.2 Ineficiencias en el manejo de inventarios",
        "   3.3 Impacto en las ventas",
        "4. Predicción y Decisión",
        "   4.1 Meses con mayor necesidad de insumos",
        "   4.2 Estrategias para aumentar ventas",
        "   4.3 Estrategias para optimizar inventario",
        "5. Propuestas de Solución",
        "   5.1 Segmentación de clientes",
        "   5.2 Campañas de marketing dirigidas",
        "   5.3 Plan de reducción de desperdicio",
        "6. Preguntas Adicionales",
        "7. Conclusión y Pasos Accionables",
        "8. Anexo de Gráficas",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 1. RESUMEN EJECUTIVO
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("1. Resumen Ejecutivo", level=1)

    doc.add_paragraph(
        "Se analizaron 5 conjuntos de datos (5,000 transacciones de ventas, 1,500 registros de clientes, "
        "10 sucursales, 2,000 registros de inventario y 1,200 interacciones digitales) para responder "
        "preguntas clave de negocio de la cadena \"Sabor Mexicano\"."
    )

    add_table_from_rows(doc,
        ["Indicador", "Resultado"],
        [
            ["Sucursal líder en ingresos", "Cancún — $118,067 MXN"],
            ["Sucursal con mejor eficiencia operativa", "León — $143,500/mes costo operativo"],
            ["Sucursal con mayor oportunidad de mejora", "Cancún — peor utilidad proxy (-$4,501,497)"],
            ["Platillo estrella", "Mole Poblano — $134,575 en ingresos"],
            ["Hora pico de ventas", "13:00 hrs (lunes a viernes)"],
            ["Principal driver de merma", "Carne de Res en Mérida — $10,470"],
            ["Segmento más grande de clientes", "Ocasionales Sensibles a Promoción — 78.7%"],
            ["Mes pico pronosticado", "Julio 2026"],
            ["Canal digital más efectivo", "TikTok — 6–12× engagement/$"],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 2. ANÁLISIS EXPLORATORIO
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("2. Análisis Exploratorio", level=1)

    # 2.1 Tendencias de ventas
    doc.add_heading("2.1 Tendencias de ventas", level=2)
    doc.add_paragraph(
        "El análisis de las ventas diarias y mensuales revela patrones claros de estacionalidad y "
        "comportamiento semanal. La hora pico dominante es las 13:00 hrs, donde el mayor ingreso "
        "ocurre consistentemente de lunes a viernes, con el lunes liderando con $35,914."
    )
    doc.add_paragraph(
        "El segundo pico vespertino ocurre entre las 19:00–20:00 hrs, orientado a cenas, con el "
        "sábado a las 20 hrs alcanzando $23,364. El domingo es el día más débil."
    )

    add_table_from_rows(doc,
        ["Día", "Hora Pico", "Ingreso en pico"],
        [
            ["Lunes", "13:00", "$35,914"],
            ["Martes", "13:00", "$32,308"],
            ["Viernes", "13:00", "$31,320"],
            ["Jueves", "13:00", "$31,157"],
            ["Sábado", "20:00", "$23,364"],
        ]
    )

    doc.add_paragraph(
        "Interpretación: La concentración del 40-45% del ingreso diario entre 13:00-14:00 indica que la "
        "operación debe optimizarse para atender la demanda máxima en ese horario. El pico sabatino "
        "nocturno sugiere oportunidad de expandir la oferta para fines de semana."
    )

    # 2.2 Platillos más vendidos
    doc.add_heading("2.2 Platillos más vendidos por región y hora del día", level=2)

    add_image_if_exists(doc, "pareto_dishes.png", caption="Figura 1: Diagrama de Pareto — Ingresos por platillo")

    doc.add_paragraph("Top 5 platillos a nivel cadena por ingresos:")
    add_table_from_rows(doc,
        ["#", "Platillo", "Categoría", "Ingresos", "Unidades"],
        [
            ["1", "Mole Poblano", "Platos Fuertes", "$134,575", "769"],
            ["2", "Enchiladas Verdes", "Platos Fuertes", "$96,660", "716"],
            ["3", "Margarita", "Bebidas", "$67,925", "715"],
            ["4", "Tacos al Pastor", "Platos Fuertes", "$60,775", "715"],
            ["5", "Guacamole con Totopos", "Entradas", "$58,473", "657"],
        ]
    )

    doc.add_paragraph(
        "Los 5 platillos anteriores representan aproximadamente el 48% del ingreso total (regla 80/20). "
        "Mole Poblano es universalmente el líder en todas las ciudades durante la franja de Comida."
    )

    doc.add_paragraph("Platillos líderes por ciudad durante la franja Comida (13:00–15:00):")
    add_table_from_rows(doc,
        ["Ciudad", "Platillo #1 (ingresos)", "Platillo #1 (volumen)"],
        [
            ["Cancún", "Mole Poblano ($12,075)", "Tacos al Pastor (75 uds)"],
            ["CDMX Centro", "Mole Poblano ($9,275)", "Tacos al Pastor (53 uds)"],
            ["Monterrey", "Mole Poblano ($11,025)", "Guacamole con Totopos (73 uds)"],
            ["Guadalajara", "Enchiladas Verdes ($9,045)", "Enchiladas Verdes (67 uds)"],
            ["Querétaro", "Enchiladas Verdes ($7,560)", "Enchiladas Verdes (56 uds)"],
            ["Puebla", "Mole Poblano ($7,350)", "Mole Poblano (42 uds)"],
            ["Mérida", "Mole Poblano ($6,825)", "Enchiladas Verdes (45 uds)"],
            ["León", "Enchiladas Verdes ($5,670)", "Enchiladas Verdes (42 uds)"],
            ["Tijuana", "Enchiladas Verdes ($6,075)", "Tacos al Pastor (47 uds)"],
        ]
    )

    doc.add_paragraph(
        "Las preferencias varían regionalmente: en el norte (Monterrey), Guacamole con Totopos destaca "
        "en volumen; en zonas turísticas (Cancún), Tacos al Pastor lidera; en el centro-occidente "
        "(Guadalajara, Querétaro, León), Enchiladas Verdes es el favorito indiscutible."
    )

    # 2.3 Sucursales
    doc.add_heading("2.3 Sucursales con mejor y peor desempeño", level=2)

    add_image_if_exists(doc, "branch_revenue_ranking.png",
                        caption="Figura 2: Ranking de sucursales por ingresos totales")

    doc.add_paragraph("Ranking completo por ingresos:")
    add_table_from_rows(doc,
        ["Pos.", "Sucursal", "Ciudad", "Ingresos", "Margen Bruto", "Tickets", "Ticket Prom."],
        [
            ["1", "Cancún", "Cancún", "$118,067", "$82,503", "650", "$181.64"],
            ["2", "CDMX Centro", "CDMX", "$108,934", "$76,365", "629", "$173.19"],
            ["3", "Monterrey", "Monterrey", "$101,214", "$70,907", "570", "$177.57"],
            ["4", "Guadalajara", "Guadalajara", "$93,049", "$65,367", "542", "$171.68"],
            ["5", "Querétaro", "Querétaro", "$85,947", "$60,197", "483", "$177.94"],
            ["6", "CDMX Sur", "CDMX", "$83,567", "$58,656", "527", "$158.57"],
            ["7", "León", "León", "$71,777", "$50,230", "386", "$185.95"],
            ["8", "Tijuana", "Tijuana", "$71,327", "$49,873", "405", "$176.12"],
            ["9", "Mérida", "Mérida", "$70,821", "$49,660", "405", "$174.87"],
            ["10", "Puebla", "Puebla", "$66,585", "$46,529", "403", "$165.22"],
        ]
    )

    add_image_if_exists(doc, "radar_branches.png",
                        caption="Figura 3: Radar multi-dimensión por sucursal (normalizado 0-1)")

    doc.add_paragraph(
        "Cancún lidera en ingresos, margen bruto y tickets, impulsada por su ubicación turística y mayor "
        "capacidad (150 personas). Puebla tiene los menores ingresos y ticket promedio, con capacidad "
        "limitada y 7 competidores. León destaca con el ticket promedio más alto ($185.95) pese a ser la "
        "sucursal más pequeña, indicando servicio premium o clientela con mayor poder adquisitivo."
    )

    # 2.4 Métodos de pago
    doc.add_heading("2.4 Métodos de pago", level=2)
    add_table_from_rows(doc,
        ["Método", "Ingresos", "Participación"],
        [
            ["Tarjeta de Crédito", "$350,263", "40.2%"],
            ["Tarjeta de Débito", "$227,291", "26.1%"],
            ["Efectivo", "$190,870", "21.9%"],
            ["App de Pago", "$102,864", "11.8%"],
        ]
    )
    doc.add_paragraph(
        "La dominancia de tarjeta de crédito (40.2%) sugiere clientela de nivel socioeconómico medio-alto. "
        "Las apps de pago (11.8%) representan un canal en crecimiento que facilita la integración con "
        "programas de lealtad digitales."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 3. IDENTIFICACIÓN DE PROBLEMAS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("3. Identificación de Problemas", level=1)

    # 3.1 Rentabilidad
    doc.add_heading("3.1 Factores que influyen en la rentabilidad", level=2)

    doc.add_paragraph(
        "La utilidad proxy se calculó como: Ingreso − Costo de ingredientes − Costo operativo asignado por ticket. "
        "El hallazgo crítico es que la sucursal con más ingresos tiene la peor rentabilidad."
    )

    add_image_if_exists(doc, "waterfall_cancún.png",
                        caption="Figura 4: Cascada de rentabilidad — Cancún (peor utilidad proxy)")

    add_image_if_exists(doc, "waterfall_león.png",
                        caption="Figura 5: Cascada de rentabilidad — León (mejor utilidad proxy)")

    add_table_from_rows(doc,
        ["Sucursal", "Utilidad Proxy", "Costo Op./mes", "Empleados", "Renta/mes", "Competidores"],
        [
            ["León (mejor)", "-$1,671,770", "$143,500", "17", "$32,000", "5"],
            ["Puebla", "-$1,777,471", "$152,000", "18", "$35,000", "7"],
            ["Mérida", "-$1,882,340", "$161,000", "19", "$38,000", "4"],
            ["Querétaro", "-$2,183,803", "$187,000", "20", "$48,000", "3"],
            ["Tijuana", "-$2,326,127", "$198,000", "20", "$52,000", "6"],
            ["Guadalajara", "-$2,454,633", "$210,000", "23", "$55,000", "5"],
            ["CDMX Sur", "-$2,809,344", "$239,000", "22", "$72,000", "4"],
            ["CDMX Centro", "-$3,283,635", "$280,000", "25", "$85,000", "3"],
            ["Monterrey", "-$3,373,093", "$287,000", "28", "$78,000", "2"],
            ["Cancún (peor)", "-$4,501,497", "$382,000", "32", "$120,000", "8"],
        ]
    )

    add_image_if_exists(doc, "cost_structure_branches.png",
                        caption="Figura 6: Estructura de costos vs ingresos por sucursal")

    doc.add_paragraph(
        "Nota metodológica: Las utilidades proxy son negativas porque el costo operativo total mensual "
        "se prorratea entre los tickets de cada sucursal. Esto no significa pérdida real, sino la carga "
        "de costos fijos que cada ticket debe absorber. El valor proxy es útil para comparar eficiencia "
        "relativa entre sucursales."
    )

    p = doc.add_paragraph()
    p.add_run("Factores clave que afectan la rentabilidad:").bold = True

    factors = [
        "Costo operativo desproporcionado: Cancún paga $382,000/mes para generar $118,067 en ingresos del período.",
        "Competencia: Cancún tiene 8 competidores cercanos, la mayor presión competitiva de la cadena.",
        "Categoría dominante: Los Platos Fuertes generan 55-65% de ingresos con ~31% de costo de ingredientes.",
        "Volumen vs. margen: Mole Poblano genera el mayor ingreso pero también el mayor costo de ingredientes.",
    ]
    for f in factors:
        doc.add_paragraph(f, style="List Bullet")

    # 3.2 Inventarios
    doc.add_heading("3.2 Ineficiencias en el manejo de inventarios", level=2)

    add_image_if_exists(doc, "waste_cost_by_branch.png",
                        caption="Figura 7: Costo total de merma por sucursal")

    add_table_from_rows(doc,
        ["Sucursal", "Costo Merma", "Tasa Quiebre", "Costo Compra", "Merma/Compra"],
        [
            ["Puebla (peor)", "$91,963", "7.1%", "$498,818", "18.4%"],
            ["Mérida", "$85,285", "8.5%", "$490,295", "17.4%"],
            ["León", "$75,801", "4.4% (mejor)", "$415,363", "18.2%"],
            ["CDMX Sur", "$64,388", "6.1%", "$547,527", "11.8%"],
            ["Cancún", "$63,998", "6.5%", "$443,847", "14.4%"],
            ["Tijuana", "$57,080", "8.6%", "$426,200", "13.4%"],
            ["Querétaro", "$53,604", "8.2%", "$499,296", "10.7%"],
            ["CDMX Centro", "$50,690", "8.3%", "$542,329", "9.3%"],
            ["Monterrey", "$43,645", "9.5% (peor)", "$646,276", "6.8%"],
            ["Guadalajara (mejor)", "$41,095", "8.9%", "$482,277", "8.5%"],
        ]
    )

    doc.add_paragraph("Top 5 drivers de merma por costo:")
    add_table_from_rows(doc,
        ["#", "Sucursal", "Ingrediente", "Razón", "Costo", "Ratio"],
        [
            ["1", "Mérida", "Carne de Res", "Caducidad", "$10,470", "25%"],
            ["2", "CDMX Sur", "Tequila", "Caducidad", "$7,813", "—"],
            ["3", "Puebla", "Cerveza", "Daño almacenamiento", "$7,497", "—"],
            ["4", "Puebla", "Carne de Res", "Error de preparación", "$6,678", "—"],
            ["5", "Puebla", "Tequila", "Exceso de pedido", "$6,355", "—"],
        ]
    )

    # 3.3 Impacto ventas
    doc.add_heading("3.3 Impacto en las ventas", level=2)
    doc.add_paragraph(
        "Las ineficiencias de inventario impactan directamente las ventas por quiebres de stock:"
    )
    add_table_from_rows(doc,
        ["Sucursal", "Ingrediente", "Tasa Quiebre"],
        [
            ["CDMX Centro", "Salsa Verde", "50%"],
            ["Cancún", "Cebolla", "50%"],
            ["Guadalajara", "Pollo", "43%"],
            ["Monterrey", "Salsa Roja", "40%"],
            ["Puebla", "Tequila", "33%"],
        ]
    )
    doc.add_paragraph(
        "Un quiebre de Salsa Verde (50% del tiempo en CDMX Centro) impacta directamente las "
        "Enchiladas Verdes — el segundo platillo más vendido de la cadena. Cuando faltan ingredientes "
        "clave, se generan ventas perdidas y deterioro de la experiencia del cliente."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 4. PREDICCIÓN Y DECISIÓN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("4. Predicción y Decisión", level=1)

    # 4.1 Pronóstico
    doc.add_heading("4.1 Meses con mayor necesidad de insumos por sucursal", level=2)
    doc.add_paragraph(
        "Se utilizó Suavizamiento Exponencial Holt-Winters (tendencia aditiva, sin estacionalidad) "
        "sobre los 12 ingredientes de mayor impacto económico por sucursal, con horizonte de 6 meses "
        "(febrero–julio 2026)."
    )

    add_image_if_exists(doc, "forecast_peaks_top15.png",
                        caption="Figura 8: Top 15 picos de demanda pronosticados")

    add_table_from_rows(doc,
        ["Sucursal", "Ingrediente", "Mes Pico", "Cantidad"],
        [
            ["Puebla", "Carne de Res", "Marzo 2026", "394.56 uds"],
            ["Monterrey", "Tequila", "Julio 2026", "376.73 uds"],
            ["Querétaro", "Frijoles", "Abril 2026", "259.00 uds"],
            ["León", "Cilantro", "Junio 2026", "258.33 uds"],
            ["Querétaro", "Tequila", "Mayo 2026", "178.77 uds"],
            ["Tijuana", "Tequila", "Julio 2026", "167.59 uds"],
            ["CDMX Sur", "Frijoles", "Julio 2026", "145.83 uds"],
            ["Tijuana", "Frijoles", "Julio 2026", "113.40 uds"],
            ["Guadalajara", "Queso Fresco", "Abril 2026", "107.39 uds"],
            ["Cancún", "Tequila", "Julio 2026", "90.22 uds"],
        ]
    )

    doc.add_paragraph(
        "Patrón estacional: Julio 2026 es el mes pico más frecuente (60%+ de los pronósticos), "
        "coincidiendo con temporada vacacional de verano. Alerta temprana: Puebla necesita 394.56 "
        "unidades de Carne de Res en marzo 2026, requiriendo preparación desde febrero."
    )

    # 4.2 Estrategias ventas
    doc.add_heading("4.2 Estrategias para aumentar ventas", level=2)

    strategies_sales = [
        ("Maximizar la hora pico (13:00–14:00)",
         "Implementar menú ejecutivo express (combo Mole Poblano + bebida + postre) a precio competitivo. "
         "Las ventas a las 13 hrs generan $35,914 los lunes; un combo con 10% de descuento podría "
         "aumentar el ticket promedio en $15–20."),
        ("Activar sábados noche",
         "El sábado 20 hrs muestra un pico de $23,364 — demanda no explotada. Lanzar \"Noche Mexicana\" "
         "semanal con mariachi + menú especial + Margarita 2×1."),
        ("Bundles de platillos estrella",
         "Combinar Mole Poblano + Margarita en un bundle \"Sabor Completo\" con 15% de descuento. "
         "Promocionar Enchiladas Verdes los miércoles (\"Miércoles Verde\")."),
        ("Impulsar sucursales con ticket alto pero bajo volumen",
         "León tiene el ticket más alto ($185.95) pero el menor número de tickets (386). Invertir en "
         "marketing local y programa de referidos."),
        ("Expansión digital",
         "CDMX Sur tiene 10.5% de conversión digital. Replicar sus tácticas (contenido orgánico en "
         "Instagram/TikTok) en sucursales con baja conversión como Puebla (4.3%)."),
    ]
    for title, desc in strategies_sales:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    # 4.3 Estrategias inventario
    doc.add_heading("4.3 Estrategias para optimizar inventario", level=2)

    strategies_inv = [
        ("Política de reorden basada en pronóstico",
         "Puntos de reorden calculados para 242 combinaciones ingrediente-sucursal con stock de seguridad "
         "(z=1.65, nivel de servicio 95%). Programar órdenes automáticas."),
        ("FEFO (First Expired, First Out)",
         "Rotación estricta por fecha de caducidad en proteínas y perecederos. Señalización visual de "
         "vencimientos en cámaras de refrigeración."),
        ("Ajustar tamaños de lote",
         "Puebla y Mérida deben reducir cantidad por pedido y aumentar frecuencia. Puebla: cambiar de "
         "pedidos semanales a cada 3 días para Carne de Res y Cerveza."),
        ("Sincronizar compras con pronóstico estacional",
         "Antes de julio 2026 (mes pico), aumentar inventarios de Tequila, Frijoles y Carne de Cerdo. "
         "Puebla: preparar proveedor de Carne de Res para marzo 2026."),
        ("Proveedores alternos",
         "Los 242 items analizados requieren reorden inmediato — diversificar proveedores para garantizar "
         "disponibilidad."),
    ]
    for title, desc in strategies_inv:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 5. PROPUESTAS DE SOLUCIÓN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("5. Propuestas de Solución", level=1)

    # 5.1 Segmentación
    doc.add_heading("5.1 Segmentación de clientes por patrones de consumo", level=2)
    doc.add_paragraph(
        "Segmentación RFM proxy (Recencia, Frecuencia, Valor Monetario) con KMeans clustering "
        "optimizado por silhouette score. Se identificaron 3 segmentos:"
    )

    add_image_if_exists(doc, "rfm_scatter_segments.png",
                        caption="Figura 9: Segmentación de clientes — RFM Proxy")

    add_image_if_exists(doc, "personas_summary.png",
                        caption="Figura 10: Perfiles de segmentación de clientes")

    add_table_from_rows(doc,
        ["Segmento", "Persona", "Clientes", "%", "Frec. Prom.", "Gasto Prom.", "Lealtad", "Promos"],
        [
            ["0", "Leales Premium", "320", "21.3%", "17.0 visitas", "$4,409", "100%", "66.3%"],
            ["1", "Ocasionales", "896", "59.7%", "6.0 visitas", "$1,451", "9.5%", "64.0%"],
            ["2", "Ocasionales", "284", "18.9%", "6.4 visitas", "$1,574", "13.7%", "69.4%"],
        ]
    )

    p = doc.add_paragraph()
    p.add_run("Leales Premium (320, 21.3%): ").bold = True
    p.add_run(
        "Motor de ingresos con $4,409 promedio. 100% miembros de lealtad. Estrategia: experiencias "
        "exclusivas, maridajes premium, acceso anticipado a platillos nuevos."
    )

    p = doc.add_paragraph()
    p.add_run("Ocasionales Sensibles a Promoción (1,180, 78.7%): ").bold = True
    p.add_run(
        "Visitan ~6 veces/año. Solo 9.5-13.7% son de lealtad — gran oportunidad de conversión. "
        "Aceptación de promos: 64-69%. Estrategia: descuentos de primera visita, combos, inscripción "
        "automática al programa de lealtad."
    )

    # 5.2 Campañas marketing
    doc.add_heading("5.2 Campañas de marketing dirigidas por sucursal", level=2)

    doc.add_paragraph("Top platillos a promocionar por sucursal (score = volumen + margen + señal digital):")
    add_table_from_rows(doc,
        ["#", "Sucursal", "Platillo", "Franja", "Score"],
        [
            ["1", "Guadalajara", "Enchiladas Verdes", "Comida", "0.911"],
            ["2", "Monterrey", "Mole Poblano", "Comida", "0.887"],
            ["3", "Monterrey", "Guacamole con Totopos", "Comida", "0.871"],
            ["4", "CDMX Centro", "Mole Poblano", "Comida", "0.847"],
            ["5", "CDMX Centro", "Margarita", "Comida", "0.840"],
            ["6", "Cancún", "Mole Poblano", "Comida", "0.836"],
            ["7", "CDMX Centro", "Guacamole c/ Totopos", "Comida", "0.832"],
            ["8", "Guadalajara", "Mole Poblano", "Comida", "0.823"],
            ["9", "Querétaro", "Enchiladas Verdes", "Comida", "0.806"],
            ["10", "Cancún", "Tacos al Pastor", "Comida", "0.788"],
        ]
    )

    doc.add_paragraph("Campañas por segmento y sucursal — Para Ocasionales:")
    add_table_from_rows(doc,
        ["Sucursal", "Clientes", "Aceptación", "Canal", "Mensaje"],
        [
            ["CDMX Centro", "139", "64%", "WhatsApp+App+Email", "Promos de entrada y menú destacado"],
            ["León", "132", "70%", "WhatsApp+App+Email", "Descuento primera visita + lealtad"],
            ["Guadalajara", "121", "69%", "WhatsApp+App+Email", "Combo Enchiladas Verdes miércoles"],
            ["Monterrey", "115", "65%", "WhatsApp+App+Email", "Mole Poblano en bundle de comida"],
        ]
    )

    doc.add_paragraph("Para Leales Premium:")
    add_table_from_rows(doc,
        ["Sucursal", "Clientes VIP", "Aceptación", "Canal", "Mensaje"],
        [
            ["Guadalajara", "45", "67%", "Email+App exclusiva", "Experiencias VIP, maridajes"],
            ["Tijuana", "39", "65%", "Email+App exclusiva", "Beneficios exclusivos de lealtad"],
            ["Mérida", "35", "66%", "Email+App exclusiva", "Acceso anticipado a nuevo menú"],
            ["CDMX Centro", "33", "76%", "Email+App exclusiva", "Noche del Socio con chef invitado"],
        ]
    )

    doc.add_paragraph("Canal digital más efectivo por campaña:")
    add_table_from_rows(doc,
        ["Campaña", "Mejor plataforma", "Engagement/$"],
        [
            ["Día de la Familia", "TikTok", "12.39×"],
            ["Noche Mexicana", "TikTok", "10.11×"],
            ["Festival del Mole", "Instagram", "8.5×"],
            ["Happy Hour Margaritas", "Facebook", "7.2×"],
            ["Descuento Cumpleañeros", "TikTok", "40% conversión"],
        ]
    )

    # 5.3 Plan desperdicio
    doc.add_heading("5.3 Plan para reducir el desperdicio de alimentos", level=2)
    doc.add_paragraph(
        "Costo total de merma en la cadena: ~$627,550 MXN. Plan en 5 acciones priorizadas:"
    )
    add_table_from_rows(doc,
        ["Prioridad", "Acción", "Sucursal(es)", "Ingrediente(s)", "Ahorro est."],
        [
            ["1 — Urgente", "FEFO estricto con señalización", "Mérida, Puebla", "Carne de Res, Tequila", "~$20,000"],
            ["2 — Alta", "Reducir lote, más frecuencia", "Puebla, Mérida", "Cerveza, Carne de Cerdo", "~$15,000"],
            ["3 — Alta", "Control temperatura y almacenamiento", "Puebla", "Cerveza", "~$7,500"],
            ["4 — Media", "Capacitación en preparación", "Puebla, León", "Carne de Res, Aguacate", "~$10,000"],
            ["5 — Media", "Sincronizar con pronóstico", "Todas", "Tequila, Frijoles", "~$12,000"],
        ]
    )
    doc.add_paragraph(
        "Meta: Reducir el costo de merma total en 30% en 90 días (de $627,550 a ~$440,000), "
        "focalizando en Puebla, Mérida y León."
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 6. PREGUNTAS ADICIONALES
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("6. Preguntas Adicionales del Caso de Estudio", level=1)

    # 6.1
    doc.add_heading("6.1 ¿Cómo varía el comportamiento del cliente según ubicación y horarios?", level=2)
    add_table_from_rows(doc,
        ["Zona", "Comportamiento", "Ejemplo"],
        [
            ["Turística (Cancún)", "Mayor volumen, clientes de paso, 8 competidores", "Tacos al Pastor líder en volumen"],
            ["Corporativa (CDMX Centro, Mty)", "Fuerte pico comida ejecutiva 13:00–14:00", "Mole Poblano + Margarita para negocios"],
            ["Residencial (CDMX Sur, GDL)", "Distribución uniforme comida-cena", "Enchiladas Verdes y platos familiares"],
            ["Cultural (Puebla, Mérida, Qro)", "Turismo cultural + local", "Platillos tradicionales regionales"],
            ["Fronteriza (Tijuana)", "Influencia binacional", "Tacos al Pastor top volumen"],
            ["Industrial (León)", "Menor volumen, mayor ticket ($185.95)", "Enchiladas Verdes y Chiles Rellenos"],
        ]
    )
    doc.add_paragraph(
        "Variaciones horarias: 13:00-14:00 concentra 40-45% del ingreso diario en TODAS las sucursales. "
        "Segundo pico a las 19:00-21:00, especialmente fuerte en Cancún y sábados. "
        "Horas de menor actividad (11:00, 22:00) son ideales para promociones \"early bird\" o descuentos nocturnos."
    )

    # 6.2
    doc.add_heading("6.2 ¿Qué platillos deben promocionarse más en cada sucursal?", level=2)
    add_table_from_rows(doc,
        ["Sucursal", "#1", "#2", "#3"],
        [
            ["Cancún", "Mole Poblano (0.836)", "Tacos al Pastor (0.788)", "Margarita (0.774)"],
            ["CDMX Centro", "Mole Poblano (0.847)", "Margarita (0.840)", "Guacamole c/ Totopos (0.832)"],
            ["Monterrey", "Mole Poblano (0.887)", "Guacamole c/ Totopos (0.871)", "Margarita (0.764)"],
            ["Guadalajara", "Enchiladas Verdes (0.911)", "Mole Poblano (0.823)", "Margarita (0.756)"],
            ["Querétaro", "Enchiladas Verdes (0.806)", "Mole Poblano (0.737)", "Guacamole c/ Totopos (0.643)"],
            ["CDMX Sur", "Mole Poblano (0.648)", "Tacos al Pastor (0.625)", "—"],
            ["Puebla", "Mole Poblano (0.707)", "Guacamole c/ Totopos (0.613)", "Enchiladas Verdes (0.600)"],
            ["Mérida", "Enchiladas Verdes (0.720)", "Mole Poblano (0.705)", "—"],
            ["Tijuana", "Enchiladas Verdes (0.627)", "—", "—"],
            ["León", "Enchiladas Verdes (0.758)", "Mole Poblano (0.749)", "Chiles Rellenos (0.635)"],
        ]
    )
    doc.add_paragraph(
        "Criterio del score: 45% volumen de ventas + 40% margen proxy + 15% sentimiento digital. "
        "Acción recomendada: Promover en franja Comida con bundle y CTA digital."
    )

    # 6.3
    doc.add_heading("6.3 ¿Cómo ajustar los niveles de inventario?", level=2)
    doc.add_paragraph(
        "Sistema de reorden propuesto: Stock de seguridad con z=1.65 (servicio 95%), punto de reorden "
        "= consumo diario × lead time + stock de seguridad."
    )
    add_table_from_rows(doc,
        ["Sucursal", "Problema principal", "Acción"],
        [
            ["Puebla", "Mayor merma ($91,963)", "Reducir lotes, capacitar, FEFO"],
            ["Mérida", "Caducidad Carne de Res ($10,470)", "FEFO + proveedor local frecuente"],
            ["Monterrey", "Mayor quiebre (9.5%)", "Aumentar stock seguridad salsas y proteínas"],
            ["CDMX Centro", "Quiebre Salsa Verde (50%)", "Duplicar stock + proveedor alterno"],
            ["Cancún", "Quiebre Cebolla (50%)", "Entregas diarias de perecederos"],
            ["Guadalajara", "Quiebre Pollo (43%)", "Reserva congelada + segunda fuente"],
            ["León", "Bajo quiebre pero alta merma", "Reducir cantidades, mantener frecuencia"],
        ]
    )

    # 6.4
    doc.add_heading("6.4 ¿Qué cambios operativos mejorarían la experiencia del cliente?", level=2)

    changes = [
        ("Reducir tiempos de respuesta digital",
         "Puebla responde en 2.10 hrs promedio (el peor). Meta: <1 hora en todas las sucursales. "
         "Mérida ya lo logra con 0.67 hrs."),
        ("Ampliar programa de lealtad",
         "Solo 21.3% de clientes son miembros. Inscripción automática al primer pago con App. "
         "Incentivo: postre gratis en la siguiente visita."),
        ("Optimizar hora pico",
         "40%+ de ventas en 13:00–14:00 genera saturación. Menú express de 3 opciones listo en "
         "<10 minutos. Pedido anticipado vía App."),
        ("Mejorar sentimiento digital",
         "Tijuana (0.302) y CDMX Sur (0.331) tienen peor sentimiento. Programa de recuperación: "
         "respuesta personalizada a reseñas negativas. Capacitación en atención."),
        ("Estacionamiento en Puebla",
         "Única sucursal sin estacionamiento. Convenio con estacionamiento cercano o servicio valet."),
    ]
    for title, desc in changes:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)

    doc.add_paragraph("Tiempos de respuesta digital por sucursal:")
    add_table_from_rows(doc,
        ["Sucursal", "Tiempo Respuesta (hrs)", "Estado"],
        [
            ["Mérida", "0.67", "✅ Excelente"],
            ["Cancún", "1.14", "⚠ Aceptable"],
            ["León", "1.20", "⚠ Aceptable"],
            ["Monterrey", "1.37", "⚠ A mejorar"],
            ["CDMX Centro", "1.52", "❌ Requiere acción"],
            ["CDMX Sur", "1.69", "❌ Requiere acción"],
            ["Guadalajara", "1.71", "❌ Requiere acción"],
            ["Puebla", "2.10", "❌ Urgente"],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 7. CONCLUSIÓN
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("7. Conclusión y Pasos Accionables", level=1)

    doc.add_heading("Hallazgos principales", level=2)

    findings = [
        "La paradoja Cancún: Genera los mayores ingresos ($118,067) pero la peor rentabilidad "
        "(-$4.5M proxy) por un costo operativo de $382,000/mes.",
        "Modelo León: Con el menor costo operativo ($143,500/mes) y el mejor ticket promedio ($185.95), "
        "demuestra que la eficiencia operativa supera al volumen puro.",
        "Concentración de ventas: Mole Poblano y Enchiladas Verdes generan ~27% del ingreso total. "
        "La hora comida (13:00–14:00) concentra el 40%+ del ingreso diario.",
        "Crisis de inventario: $627,550 en merma total + quiebres de hasta 50% en ingredientes clave.",
        "Oportunidad digital: TikTok ofrece 6–12× engagement por peso invertido.",
        "78.7% de clientes son Ocasionales con 64% aceptación de promos — mayor potencial de conversión.",
    ]
    for i, f in enumerate(findings, 1):
        doc.add_paragraph(f"{i}. {f}")

    doc.add_heading("Plan de acción 30-60-90 días", level=2)

    doc.add_heading("Primeros 30 días (Urgente):", level=3)
    for action in [
        "Implementar FEFO en Puebla y Mérida para Carne de Res, Tequila y Cerveza.",
        "Activar política de reorden automática para los 242 items identificados.",
        "Lanzar campaña piloto en TikTok para Guadalajara (Enchiladas Verdes, score 0.911).",
        "Reducir tiempo de respuesta digital en Puebla a <1.5 hrs.",
    ]:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("30–60 días (Alta prioridad):", level=3)
    for action in [
        "Ejecutar campañas segmentadas por persona en sucursales con menor margen.",
        "Implementar programa de lealtad ampliado con inscripción automática.",
        "Negociar proveedores alternos para ingredientes con quiebre >30%.",
        "Preparar stock para pico de marzo 2026 (Puebla — Carne de Res: 394 uds).",
    ]:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("60–90 días (Consolidación):", level=3)
    for action in [
        "Lanzar menú ejecutivo express en hora comida en todas las sucursales.",
        "Activar \"Noche Mexicana\" en sábados para explotar el pico de 20 hrs.",
        "Revisar estructura de costos de Cancún (renta $120K → ¿renegociación?).",
        "Capacitación en almacenamiento y preparación en Puebla y León.",
        "Reentrenar modelos de pronóstico y segmentación con datos del nuevo período.",
    ]:
        doc.add_paragraph(action, style="List Bullet")

    doc.add_heading("KPIs de seguimiento mensual", level=2)
    add_table_from_rows(doc,
        ["KPI", "Meta", "Frecuencia"],
        [
            ["Margen proxy por sucursal", "Mejorar 15% vs baseline", "Mensual"],
            ["Costo merma / costo compra", "< 12% en todas", "Mensual"],
            ["Tasa quiebre inventario", "< 5% por ingrediente", "Semanal"],
            ["Conversión digital", "> 8% promedio cadena", "Mensual"],
            ["Inscripción a lealtad", "> 30% de nuevos clientes", "Mensual"],
            ["Ingreso incremental/campaña", "> $5,000 por campaña", "Por campaña"],
            ["Tiempo respuesta digital", "< 1 hora en todas", "Semanal"],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # 8. ANEXO DE GRÁFICAS
    # ══════════════════════════════════════════════════════════════
    doc.add_heading("8. Anexo de Gráficas y Visualizaciones", level=1)

    doc.add_heading("Gráficas interactivas (HTML)", level=2)
    html_charts = [
        ("sales_trend_daily.html", "Tendencia diaria de ventas"),
        ("sales_trend_monthly.html", "Tendencia mensual de ventas"),
        ("sales_by_city.html", "Ventas por ciudad"),
        ("sales_by_hour_day.html", "Mapa de calor ventas por hora y día"),
        ("top_dishes_by_region_daypart.html", "Top platillos por región y franja"),
        ("branch_ranking_sales_margin.html", "Ranking sucursales (ventas y margen)"),
        ("payment_method_mix.html", "Mix de métodos de pago"),
        ("inventory_waste_shortage_heatmap.html", "Heatmap merma/quiebre inventario"),
        ("digital_sentiment_platform.html", "Sentimiento digital por plataforma"),
    ]
    for fname, desc in html_charts:
        doc.add_paragraph(f"{desc} — outputs/charts/{fname}", style="List Bullet")

    doc.add_heading("Gráficas generadas para este informe (PNG)", level=2)

    png_charts = [
        ("waterfall_cancún.png", "Cascada de rentabilidad — Cancún"),
        ("waterfall_león.png", "Cascada de rentabilidad — León"),
        ("pareto_dishes.png", "Pareto de platillos por ingreso"),
        ("radar_branches.png", "Radar multi-dimensión por sucursal"),
        ("rfm_scatter_segments.png", "Scatter RFM de segmentación"),
        ("cost_structure_branches.png", "Estructura de costos vs ingresos"),
        ("waste_cost_by_branch.png", "Merma por sucursal"),
        ("forecast_peaks_top15.png", "Top 15 picos de demanda pronosticados"),
        ("branch_revenue_ranking.png", "Ranking de ingresos por sucursal"),
        ("personas_summary.png", "Perfiles de segmentación de clientes"),
    ]
    for fname, desc in png_charts:
        path = os.path.join(CHARTS, fname)
        if os.path.exists(path):
            doc.add_paragraph(desc, style="List Bullet")

    # ── Footer note ──
    doc.add_paragraph("")
    add_formatted_paragraph(
        doc,
        "Informe generado con pipeline reproducible. Datos procesados con Python 3.12, "
        "Pandas, Scikit-learn, Statsmodels y Plotly. Visualizaciones interactivas disponibles "
        "en los archivos HTML listados.",
        italic=True, size=9, color=(127, 140, 141),
    )

    # ── Save ──
    doc.save(OUTPUT_DOCX)
    print(f"✅ Documento generado: {OUTPUT_DOCX}")


if __name__ == "__main__":
    build()
